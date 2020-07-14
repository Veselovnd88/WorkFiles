import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import docx.shared
import openpyxl


class WordTemplate:
    """Open template file and initializating all table's cells content
    for next activity"""

    def __init__(self, filename):
        """Конструктор, сразу инициализирует переменные"""
        self.filename = filename  # название открываемого файла
        self.doc = docx.Document(self.filename)  # открывает докх документ
        self.offer_num = self.doc.tables[1].rows[0].cells[0]  # место где находится номер ТКП
        self.customer_name = self.doc.tables[1].rows[0].cells[2]  # место, где находится имя клиента
        self.main_table = self.doc.tables[2]  # основная таблица
        self.rows = 0  # переменная - номер строки, по умолчанию ноль, как только создается строка - +1
        self.offer_head = self.doc.tables[2].rows[0]
        self.pos_head = self.offer_head.cells[0].text
        self.name_head = self.offer_head.cells[1].text
        self.qnt_head = self.offer_head.cells[2].text
        self.deltime_head = self.offer_head.cells[3].text
        self.price_head = self.offer_head.cells[4]
        self.total_head = self.offer_head.cells[5]

    def create_main(self):
        """
        Заполнение констант ТКП
        """
        head_dict = ExcelParse().header()
        """ Заполнение и центрование даты"""
        number = str(head_dict['Дата'])[8:10] + str(head_dict['Дата'])[5:7] + \
                 '-' + str(head_dict['Дата'])[2:4]
        date_of = ' от ' + str(head_dict['Дата'])[8:10] + '.' + \
                  str(head_dict['Дата'])[5:7] + '.' + \
                  str(head_dict['Дата'])[0:4] + ' г.'
        self.offer_num.text = '№ ' + head_dict['Имя ТКП'] + number + date_of
        self.aligment_cell(self.offer_num)

        """Заполнение и центрование Заказчика"""
        self.customer_name.text = head_dict['Заказчик']
        self.aligment_cell(self.customer_name)

        """ЦЕны"""
        self.price_head.text = 'Цена ' + head_dict['Цена']
        self.aligment_cell(self.price_head)
        self.total_head.text = 'Сумма ' + head_dict['Сумма']
        self.aligment_cell(self.total_head)

        """ Условия поставки"""
        self.doc.paragraphs[7].text = 'Условия оплаты:'
        self.doc.paragraphs[7].runs[0].bold = True

        self.doc.add_paragraph(head_dict['Условия оплаты'])
        self.doc.add_paragraph('Условия доставки:')
        self.doc.paragraphs[9].runs[0].bold = True
        self.doc.add_paragraph(head_dict['Условия доставки'])
        self.doc.add_paragraph('Документация:')
        self.doc.paragraphs[11].runs[0].bold = True
        self.doc.add_paragraph(head_dict['Документация'] + '\n')

        self.doc.add_paragraph('Исполнитель:')
        self.doc.paragraphs[13].runs[0].font.size = Pt(10)  # Размер шрифта
        self.doc.paragraphs[13].runs[0].bold = True  # Жирный шрифт
        self.doc.add_paragraph(head_dict['Куратор'])
        self.doc.paragraphs[14].runs[0].font.size = Pt(10)
        self.doc.add_paragraph('+7 (495) 921 30 12 доб.' + self.addnum(head_dict))
        self.doc.paragraphs[15].runs[0].font.size = Pt(10)

    @staticmethod
    def addnum(head_dict):
        """ Добавочный номер телефона в зависимости от фамилии"""
        if head_dict['Куратор'] == 'Веселов Н.Д.':
            return '1025'
        elif head_dict['Куратор'] == 'Казаков Д.В.':
            return '1026'
        elif head_dict['Куратор'] == 'Бондарцев М.А.':
            return '1027'

    @staticmethod
    def aligment_cell(field):
        """Вырванивает по центру"""
        field_paragraph = field.paragraphs[0]
        field_paragraph.text = field.text
        field_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def save(self, newfile=None):
        if newfile is None:
            self.doc.save(self.filename)
        else:
            self.doc.save(newfile)

    def generate_rows(self, params: list):
        print('Params', params)
        for i in params:
            self.main_table.add_row()
            self.rows += 1
            new_row = self.main_table.rows[self.rows]
            pos = new_row.cells[0]
            pos.text = str(i['№ поз.'])  # Номер позиции
            name = new_row.cells[1]  # Переход в ячейку с наименованием
            name_fill = i['Брэнд'] + '\n' + i['Наименование поз.'] + ' ' + i['Брэнд'] + '. ' + 'Модель ' + \
                        i['Модель'] +'. ' + \
                         i['Кодировка'] + '\n' + i['Расшифровка']
            name.text = name_fill
            qnt = i['Кол-во, шт.']
            new_row.cells[2].text = str(qnt)
            self.aligment_cell(new_row.cells[2])
            deltime = i['Срок поставки']
            new_row.cells[3].text = deltime
            self.aligment_cell(new_row.cells[3])
            price = '%0.2f' % (i['Цена (цифры)'])

            new_row.cells[4].text = str(price).replace('.', ',')
            self.aligment_cell(new_row.cells[4])
            total = '%0.2f' % i['Сумма (цифры)']
            new_row.cells[5].text = str(total).replace('.', ',')
            self.aligment_cell(new_row.cells[5])


class ExcelParse:
    """
    Class for parsing excel File, each row - parameters for the filling word table
    """

    def __init__(self):
        self.basic_cols = [1, 2, 3, 4, 5, 6, 7, 8, 9]  # Вкладка Actual
        self.main_cols = [1, 2, 3, 4, 5, 6, 7, 8, 9, 12]  # Вкладка Technic
        self.head_dict = {}
        self.main_dict = {}
        self.mainlst = []
        sheet = openpyxl.load_workbook('data.xlsx', data_only=True)
        self.offer_data = sheet["Actual"]
        self.rows_data = sheet['Technic']
        print(self.offer_data.max_column)
        self.basic_head = {}  # заголовок Актуал
        self.basic_main = {}  # заголовок Техник

    def header(self):
        """Считывание основных данных для заполнения ТКП,
        читаем колонки
        1-Дата; 2-Заказчик; 3-Имя ТКп; 4-Тип цены
        5 - тип суммы
        6- условия оплаты
        7 - условия доставки
        8-Документация
        9 - Куратор
        """
        for i in self.basic_cols:
            self.basic_head[i] = self.offer_data.cell(row=1, column=i).value  # заполнение заголовка
            self.head_dict[self.basic_head[i]] = self.offer_data.cell(row=2, column=i).value  # значения
        print(self.head_dict)

        return self.head_dict

    def rows(self):
        for k in range(self.rows_data.max_row):
            self.main_dict = {}
            if self.rows_data.cell(row=k + 2, column=1).value is None:
                break
            for i in self.main_cols:
                self.basic_main[i] = self.rows_data.cell(row=1, column=i).value
                if self.rows_data.cell(row=k + 2, column=i).value is None:
                    self.main_dict[self.basic_main[i]] = ''
                else:
                    self.main_dict[self.basic_main[i]] = self.rows_data.cell(row=k + 2, column=i).value

            #  print(self.main_dict)
            self.mainlst.append(self.main_dict)
        print(self.mainlst)
        return self.mainlst
        # return self.main_dict


def main():
    newdoc = WordTemplate('testoff.docx')
    newdoc.create_main()
    newdoc.generate_rows(ExcelParse().rows())
    newdoc.save()
    # newex = ExcelParse()
    # newex.rows()


if __name__ == '__main__':
    main()
